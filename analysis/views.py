from django.http import JsonResponse
from rest_framework import status
from rest_framework.decorators import api_view
from rest_framework.response import Response
from pandas import Series
import pandas as pd
import jieba
import docx
from docx import Document
from rest_framework.renderers import  JSONRenderer  #导入序列化数据转json数据的模块


#接收上传Excel文件方法
@api_view(['POST'])
def upload_excel(request):
    # 要加上r，否则会出现500错误(r'file')
    file = request.FILES.get(r'file')
    print(file)
    data = pd.read_excel(file)

    # 获得列标签
    columns = list(data.columns.values)
    print("输出列标签")
    print(columns)

    # print("输出to_dict方法转化字典结果")
    print(data)
    rows =data.to_dict(orient= 'records')
    # print(rows)
    json_dict = {'columns': columns,'rows': rows}
    print(json_dict)
    return JsonResponse(json_dict,status=status.HTTP_200_OK)
    # return Response(columns, status=status.HTTP_200_OK)


#接收上传word、txt文件并处理
@api_view(['POST'])
def handle_word_txt(request):
    file = request.FILES.get(r'file')
    print(file)
    file_name = str(file)
    content = '' #用于存储读取的文件文本内容
    counts = {}     # 通过键值对的形式存储词语及其出现的次数
    dict_counts = [] #用于存储拼接出现次数前20的词语数据
    print(file_name.split('.')[1])

    if file_name.split('.')[1]=='txt': # 判断是否为txt文件
        content = handle_txt(file)

    if file_name.split('.')[1]=='docx':# 判断是否为后缀docx的word文件
        content = handle_word(file)

    if file_name.split('.')[1]=='doc':# 判断是否为后缀doc的word文件
        content = handle_word(file)
    # words = jieba.cut(content,cut_all=True) # 使用全模式对文本进行分词
    words = jieba.lcut(content) # 使用精确模式对文本进行分词
    for word in words:
        if len(word) == 1:  # 单个词语不计算在内
            continue
        else:
            counts[word] = counts.get(word,0) + 1 # 遍历所有词语，每出现一次其对应的值加 1
    items = list(counts.items())
    items.sort(key=lambda x: x[1], reverse=True)    # 根据词语出现的次数进行从大到小排序
    columns = ['word', 'count']

    # 此处循环用于拼接出现次数前20的词语数据
    for i in range(20):
        word, count = items[i]
        # print("{0:<5}{1:>5}".format(word, count))
        # list转Json
        dict = {'word':word,'count':count}
        dict_counts.append(dict)

    json_dict = {'columns': columns,'rows': dict_counts} # 将数据封装为前端需要的格式

    print("输出最终回传的结果")
    print(json_dict)
    return JsonResponse(json_dict,status=status.HTTP_200_OK)




# 处理txt，提取文本内容并返回
def handle_txt(file):
    #使用error_bad_lines=False会导致结果变得很少，被忽略部分过多
    content = str(pd.read_table(file,"r",error_bad_lines=False))
    return (content)



# 处理word，提取文本内容并返回
def handle_word(file):
    file=docx.Document(file)
    print("段落数:"+str(len(file.paragraphs)))#段落数
    content = ''
    #输出每一段的内容
    for para in file.paragraphs:
        # print(para.text)
        string = str(para.text)
        content = content + string
    print("读取word后的文本内容")
    print(content)
    return (content)